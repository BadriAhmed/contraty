import logging
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.models.contract import Language, Contract

logger = logging.getLogger(__name__)


def _add_bottom_border(paragraph, color: str = "1a365d", size: str = "12"):
    """Draw a horizontal line under the paragraph via a bottom border."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


class DOCXRenderer:
    def render_contract(self, contract: Contract, language: Language) -> bytes:
        doc = Document()
        direction = "rtl" if language == Language.ar else "ltr"
        title = contract.title_ar if language == Language.ar else contract.title_fr

        # Title
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if language == Language.ar:
            heading.runs[0].font.rtl = True
        _add_bottom_border(heading)

        # Sections
        for section in contract.sections:
            sec_title = section.title_ar if language == Language.ar else section.title_fr
            if sec_title:
                h = doc.add_heading(sec_title, level=2)
                if language == Language.ar:
                    h.runs[0].font.rtl = True

            for article in section.articles:
                text = article.text_ar if language == Language.ar else article.text_fr
                if text.strip():
                    p = doc.add_paragraph(text)
                    if language == Language.ar:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in p.runs:
                            run.font.rtl = True

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


docx_renderer = DOCXRenderer()
